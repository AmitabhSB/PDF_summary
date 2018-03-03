import sys
import traceback
import os
import time
import shutil
import PyPDF2
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfinterp import resolve1
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from cStringIO import StringIO
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lsa import LsaSummarizer as Lsa
from sumy.summarizers.luhn import LuhnSummarizer as Luhn
from sumy.summarizers.text_rank import TextRankSummarizer as TextRank
from sumy.summarizers.lex_rank import LexRankSummarizer as LexRank
from sumy.summarizers.sum_basic import SumBasicSummarizer as SumBasic
from sumy.summarizers.kl import KLSummarizer as KLsum
from sumy.nlp.stemmers import Stemmer
from sumy.utils import get_stop_words
from docx import Document


##Functon to convert pdf into text 
def convert(fname, pages=None):
    infile = file(fname, 'rb')
    content=""
    parser =PDFParser(infile)
    document=PDFDocument(parser)
    # This will give you the count of pages
    count= (resolve1(document.catalog['Pages'])['Count'])
    if not pages:
        pagenums = set()
    else:
        pagenums = count
  #  Check :  print ('converting......')
    codec = 'utf-8'
    laparams = LAParams()
    output = StringIO()
    manager = PDFResourceManager()
    converter = TextConverter(manager, output,codec=codec, laparams=laparams)
    interpreter = PDFPageInterpreter(manager, converter)

    for page in PDFPage.get_pages(infile, pagenums):
        interpreter.process_page(page)
        
    infile.close()
    converter.close()
    text = output.getvalue()
    output.close
    return text

 ##Summary Algorithms
def algo(x):
    return {
        '1' : 'Luhn',
        '2' : 'Lsa',
        '3' : 'LexRank',
        '4' : 'TextRank',
        '5' : 'SumBasic',
        '6' : 'KLsum',
        '0' : 'exit',
        }[x]

 ##Text files names with suffix 
def pdfTitle(y):
    return {
        1 : 'ABSTRACT',
        2 : 'BACKGROUND OF THE INVENTION',
        3 : 'BRIEF DESCRIPTION OF THE INVENTION',
        4 : 'BRIEF DESCRIPTION OF THE DRAWINGS',
        5 : 'DETAILED DESCRIPTION OF THE INVENTION',
        6 : 'CLAIM',
        }[y]


########################  Main Program   ########################

##Set parameters
LANGUAGE = "English"
SENTENCES_COUNT = 20

pdfDir =raw_input("Enter patent file path \n")
pdfNum= raw_input("Enter patent file name/number \n")
sourcePDFFile= pdfDir+pdfNum+'.pdf'

## Check and Verify if PDF is ready for splitting
if os.path.exists(sourcePDFFile):
    print('Found source PDF file')
elif not os.path.exists(sourcePDFFile):
    print('Source PDF not found, sleeping...')
    sys.exit(0)
    
## PDF_SummaryDir= raw_input("Enter Output Directory path \n")
chooseAlgo = algo(raw_input("Select Algorithm \n press 1 and enter for Luhn. \n press 2 and enter for Lsa. \n press 3 and enter for LexRank. \n press 4 and enter for TextRank. \n press 5 and enter for SumBasic.\n press 6 and enter for KLsum.\n press 0 and enter to exit. \n"))

## Output Directory for Text Files
outputTXTDir = os.path.dirname(pdfDir + '\Text_Files\\')
if not os.path.exists(outputTXTDir):
    os.makedirs(pdfDir + '\Text_Files\\')

## Output Directory for Summary
outputSummaryDir = os.path.dirname(pdfDir + '\Summary\\')
if not os.path.exists(outputSummaryDir):
    os.makedirs(pdfDir + '\Summary\\')

## Append backslash to output dir for txt if necessary
if not outputTXTDir.endswith('\\'):
    outputTXTDir = outputTXTDir + '\\'

## Append backslash to output dir for pdf if necessary
if not outputSummaryDir.endswith('\\'):
    outputSummaryDir = outputSummaryDir + '\\'

## Time of the creation of file
timeSuffixSummary = str(time.strftime("%d-%m-%Y_%H.%M.%S"))

        
if __name__ == '__main__':
        
        text=convert(sourcePDFFile)
        txtFileName = pdfNum +'.txt'

##Calling convert function and writing each chapter to the .txt file
        txtOutputFile = open(outputTXTDir + txtFileName, 'w')
        txtOutputFile.write(text)
        txtOutputFile.close()
        ## Opening text files and name acc to title of the file 
        with open(outputTXTDir + txtFileName, 'r') as fr:
            l=pdfTitle(1)
            fABS = open(outputTXTDir + pdfNum +l+'.txt' , 'w')
            l=pdfTitle(2)
            fBII = open(outputTXTDir + pdfNum +l+'.txt', 'w')
            l=pdfTitle(3)
            fSTI = open(outputTXTDir + pdfNum + l +'.txt', 'w')
            l=pdfTitle(4)
            fBDD = open(outputTXTDir + pdfNum + l+'.txt', 'w')
            l=pdfTitle(5)
            fDDI = open(outputTXTDir + pdfNum + l+'.txt', 'w')
            l=pdfTitle(6)
            fCLA = open(outputTXTDir + pdfNum + l+'.txt', 'w')

            copy= False
            ## Compare strings and and save line by line into text file
            ## move pointer to the begning after every section is extracted
            for line in fr:
                if line.startswith('ABSTRACT \n'):
                    copy = True
                    fABS.write(line)
                elif line.startswith('BACKGROUND OF THE INVENTION'):
                    copy =False
                elif copy:
                    fABS.write(''.join(line))
            fABS.close()
            fr.seek(0,0);
            for line in fr:
                if line.startswith('BACKGROUND OF THE INVENTION \n'):
                    copy = True
                    fBII.write(line)
                elif(line.startswith('SUMMARY OF THE INVENTION \n') or line.startswith('BRIEF DESCRIPTION OF THE INVENTION \n')):
                    copy= False
                elif copy :
                    fBII.write(''.join(line))
            fBII.close()
            fr.seek(0,0);
            for line in fr :
                if (line.startswith('SUMMARY OF THE INVENTION \n') or line.startswith('BRIEF DESCRIPTION OF THE INVENTION \n')):
                    copy = True
                    fSTI.write(line)
                elif line.startswith('BRIEF DESCRIPTION OF THE DRAWING \n'):
                    copy = False
                elif copy:
                    fSTI.write(''.join(line))
            fSTI.close()
            fr.seek(0,0);
            for line in fr :
                if line.startswith('BRIEF DESCRIPTION OF THE DRAWING \n'):
                    copy = True
                    fBDD.write(line)
                elif (line.startswith('DETAILED DESCRIPTION OF THE \n') or line.startswith('DETAILED DESCRIPTION OF THE ')):
                    copy =False
                elif copy:
                    fBDD.write(''.join(line))
            fBDD.close()
            fr.seek(0,0);
            for line in fr :
                if (line.startswith('DETAILED DESCRIPTION OF THE ') or line.startswith('DETAILED DESCRIPTION OF THE \n')):
                    copy= True
                    fDDI.write(line)
                elif line.startswith('I claim: \n'):
                    copy= False
                elif copy :
                    fDDI.write(''.join(line))
            fDDI.close()
            fr.seek(0,0);
            for line in fr:
                if line.startswith('I claim: \n'):
                    copy= True
                    fCLA.write(line)
                elif copy:
                    fCLA.write(''.join(line))
            fCLA.close()
            fr.close()
            os.unlink(outputTXTDir + txtFileName)
            
        ##  for plain text files create Summary
            doc = Document()
            doc.add_heading('Patent Summary Sheet',0).bold = True
            doc.add_heading('\nPatent Number : '+pdfNum,1).bold= True
            doc.add_heading('\nPatent Summary: ',1).bold= True
            
            for i in range(1,7):
                readFile=outputTXTDir + pdfNum +pdfTitle(i) +'.txt'
                txtOutputFile = open(readFile, 'r')
                parser = PlaintextParser.from_file(readFile , Tokenizer(LANGUAGE))
                stemmer = Stemmer(LANGUAGE)
            ##  Select from different algorithms to create summary by using different algorithms 
                if chooseAlgo == 'Lsa' :
                    summarizer = Lsa(stemmer)
                elif chooseAlgo == 'LexRank':
                    summarizer = LexRank(stemmer)
                elif  chooseAlgo == 'TextRank':
                    summarizer = TextRank(stemmer)
                elif  chooseAlgo == 'Luhn':
                    summarizer = Luhn(stemmer)
                elif  chooseAlgo == 'SumBasic':
                    summarizer = SumBasic(stemmer)
                elif  chooseAlgo == 'KLsum':
                    summarizer = KLsum(stemmer)
                else :
                    print ( 'Wrong Algorithm selected.')
                    sys.exit(0)

                stop_words = get_stop_words(LANGUAGE)
                ##  Open file in append mode so that summary will be added at the bottom of file 
                summaryOutputFile = open(outputSummaryDir + chooseAlgo + '_Summary_File'+ pdfTitle(i) + pdfNum + '.txt','a')
                p = doc.add_heading(pdfTitle(i),2).bold= True

                for sentence in summarizer(parser.document, SENTENCES_COUNT):
                    ##check : print (sentence)
                    summaryOutputFile.write(str(sentence))
                    doc.add_paragraph(unicode(str(sentence), 'utf-8'))

                txtOutputFile.close()
                summaryOutputFile.close()
            doc.add_heading('\nPatent Critique: \n',1).bold=True
            doc.add_heading('\nPatent Life remaining: \n',1).bold=True
            doc.add_heading('\nSome more useful information: \n',1).bold=True
            doc.save('Patent '+ pdfNum +'.docx')
            #os.unlink(readFile)





